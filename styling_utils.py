"""
Styling Utilities Module
Provides professional formatting and styling for Word and PDF reports
"""
import os
import logging
from typing import Optional, Tuple
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Mm, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml.shared import OxmlElement
from docx.oxml import parse_xml
from docx.oxml.shape import CT_Picture
from docx.oxml.xmlchemy import BaseOxmlElement, OneAndOnlyOne
import config

logger = logging.getLogger(__name__)


def hex_to_rgb(hex_color: str) -> RGBColor:
    """Convert hex color string to RGBColor object"""
    hex_color = hex_color.lstrip('#')
    return RGBColor(
        int(hex_color[0:2], 16),
        int(hex_color[2:4], 16),
        int(hex_color[4:6], 16)
    )


class WordStyler:
    """
    Word document styling utilities
    Applies professional formatting to Word documents
    """

    @staticmethod
    def set_font(run, font_name: str, font_size: int, bold: bool = False,
                 color: Optional[str] = None):
        """
        Set font properties for a run

        Args:
            run: Document run object
            font_name: Font name (supports Chinese fonts)
            font_size: Font size in points
            bold: Whether text is bold
            color: Hex color string (e.g., '#0066CC')
        """
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.font.bold = bold

        # Set Chinese font support
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

        if color:
            run.font.color.rgb = hex_to_rgb(color)

    @staticmethod
    def set_color(run, color_hex: str):
        """Set text color from hex string"""
        run.font.color.rgb = hex_to_rgb(color_hex)

    @staticmethod
    def set_spacing(paragraph, line_spacing: float = None,
                    before: float = None, after: float = None):
        """
        Set paragraph spacing

        Args:
            paragraph: Paragraph object
            line_spacing: Line spacing multiplier (e.g., 1.5)
            before: Space before in points
            after: Space after in points
        """
        # Set line spacing
        if line_spacing is not None:
            paragraph.paragraph_format.line_spacing = line_spacing

        # Set space before/after
        if before is not None:
            paragraph.paragraph_format.space_before = Pt(before)
        if after is not None:
            paragraph.paragraph_format.space_after = Pt(after)

    @staticmethod
    def set_first_line_indent(paragraph, chars: int = 2):
        """
        Set first line indent (Chinese characters)

        Args:
            paragraph: Paragraph object
            chars: Number of characters to indent
        """
        # Approximate character width in points (varies by font)
        char_width = Pt(config.FONT_SIZE_BODY)
        paragraph.paragraph_format.first_line_indent = char_width * chars

    @staticmethod
    def add_section_divider(doc, color: str = None, thickness_pt: float = 0.5):
        """
        Add a horizontal divider line between sections

        Args:
            doc: Document object
            color: Hex color string (defaults to COLOR_DIVIDER)
            thickness_pt: Line thickness in points
        """
        if color is None:
            color = config.COLOR_DIVIDER

        p = doc.add_paragraph()
        p_format = p.paragraph_format
        p_format.space_before = Pt(12)
        p_format.space_after = Pt(12)

        # Add bottom border to paragraph
        pPr = p._element.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')

        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), str(int(thickness_pt * 8)))  # Convert to eighth-points
        bottom.set(qn('w:space'), '4')
        bottom.set(qn('w:color'), color.lstrip('#'))

        pBdr.append(bottom)
        pPr.append(pBdr)

        return p

    @staticmethod
    def apply_heading_style(doc, text: str, level: int, emoji: str = '') -> None:
        """
        Apply consistent heading style with optional emoji

        Args:
            doc: Document object
            text: Heading text (without emoji)
            level: Heading level (1, 2, or 3)
            emoji: Optional emoji prefix (single character recommended)
        """
        heading = doc.add_heading(f'{emoji} {text}' if emoji else text, level=level)

        # Apply style based on level
        if level == 1:
            WordStyler.set_font(
                heading.runs[0],
                config.FONT_HEADING1,
                config.FONT_SIZE_HEADING1,
                bold=True,
                color=config.COLOR_ACCENT_BLUE
            )
        elif level == 2:
            WordStyler.set_font(
                heading.runs[0],
                config.FONT_HEADING2,
                config.FONT_SIZE_HEADING2,
                bold=True,
                color=config.COLOR_TEXT_DARK
            )
        else:  # level 3
            WordStyler.set_font(
                heading.runs[0],
                config.FONT_BODY,
                config.FONT_SIZE_BODY,
                bold=True,
                color=config.COLOR_TEXT_DARK
            )

        # Set spacing
        WordStyler.set_spacing(
            heading,
            line_spacing=config.LINE_SPACING,
            before=Pt(12),
            after=Pt(6)
        )

    @staticmethod
    def apply_body_style(doc, text: str, bold_prefix: str = '',
                         indent: bool = False) -> None:
        """
        Add a body paragraph with consistent styling

        Args:
            doc: Document object
            text: Paragraph text
            bold_prefix: Text to make bold at the start
            indent: Whether to apply first-line indent
        """
        p = doc.add_paragraph()

        if bold_prefix:
            run = p.add_run(bold_prefix)
            WordStyler.set_font(run, config.FONT_BODY, config.FONT_SIZE_BODY, bold=True)
            run = p.add_run(text)
            WordStyler.set_font(run, config.FONT_BODY, config.FONT_SIZE_BODY, bold=False)
        else:
            p.add_run(text)
            WordStyler.set_font(p.runs[0], config.FONT_BODY, config.FONT_SIZE_BODY, bold=False)

        # Set spacing
        WordStyler.set_spacing(
            p,
            line_spacing=config.LINE_SPACING,
            before=Pt(3),
            after=Pt(3)
        )

        # Apply indent if requested
        if indent:
            WordStyler.set_first_line_indent(p, config.FIRST_LINE_INDENT_CHARS)

    @staticmethod
    def set_page_margins(doc, top_cm: float = None, bottom_cm: float = None,
                         left_cm: float = None, right_cm: float = None):
        """
        Set document page margins

        Args:
            doc: Document object
            top_cm: Top margin in cm
            bottom_cm: Bottom margin in cm
            left_cm: Left margin in cm
            right_cm: Right margin in cm
        """
        sections = doc.sections
        for section in sections:
            if top_cm is not None:
                section.top_margin = Cm(top_cm)
            if bottom_cm is not None:
                section.bottom_margin = Cm(bottom_cm)
            if left_cm is not None:
                section.left_margin = Cm(left_cm)
            if right_cm is not None:
                section.right_margin = Cm(right_cm)

    @staticmethod
    def add_paragraph_shading(paragraph, fill_color: str = 'FFFFFF', opacity: float = 0.8):
        """
        Add semi-transparent shading to a paragraph

        Note: Word doesn't support true transparency for paragraph shading.
        This method applies solid shading. For transparency effect, use lighter colors.

        Args:
            paragraph: Word paragraph object
            fill_color: Shading color (hex, e.g., 'FFFFFF' for white)
            opacity: Opacity (0.0-1.0) - simulated by adjusting color brightness
        """
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        # Get or create paragraph properties
        pPr = paragraph._element.get_or_add_pPr()

        # Create shading element
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), fill_color)

        # Add to paragraph properties
        pPr.append(shd)


class BackgroundTemplateManager:
    """
    Manages background template extraction and application
    """

    def __init__(self, template_path: str = None):
        """
        Initialize background template manager

        Args:
            template_path: Path to the cat background template document
        """
        self.template_path = template_path or config.CAT_BACKGROUND_TEMPLATE
        self._background_image = None
        self._background_width = None
        self._background_height = None
        self._transparent_image_cache = {}  # Cache for processed images with different opacity

    def _load_template(self) -> Optional[bytes]:
        """
        Load and extract background image from template document

        Returns:
            Image data as bytes, or None if not found
        """
        if not os.path.exists(self.template_path):
            logger.warning(f"Background template not found: {self.template_path}")
            return None

        try:
            # Method 1: Try to find image in document relationships
            template_doc = Document(self.template_path)

            # Extract images from the document
            for rel in template_doc.part.rels.values():
                if "image" in rel.target_ref:
                    image_part = rel.target_part
                    image_data = image_part.blob

                    # Get image dimensions
                    from PIL import Image
                    import io

                    img = Image.open(io.BytesIO(image_data))
                    self._background_width, self._background_height = img.size

                    logger.info(f"Loaded background image: {self._background_width}x{self._background_height}px")
                    return image_data

            # Method 2: Try to extract from media folder (for templates with header images)
            import zipfile
            with zipfile.ZipFile(self.template_path, 'r') as zip_ref:
                media_files = [n for n in zip_ref.namelist() if 'word/media/' in n]
                if media_files:
                    # Get the first image
                    image_data = zip_ref.read(media_files[0])

                    # Get image dimensions
                    from PIL import Image
                    import io

                    img = Image.open(io.BytesIO(image_data))
                    self._background_width, self._background_height = img.size

                    logger.info(f"Loaded background image from media: {self._background_width}x{self._background_height}px")
                    return image_data

            logger.warning("No background image found in template")
            return None

        except Exception as e:
            logger.error(f"Failed to load background template: {e}")
            import traceback
            traceback.print_exc()
            return None

    def _process_image_transparency(self, image_data: bytes, opacity: float = 0.6) -> bytes:
        """
        Process image to apply transparency effect using PIL

        Args:
            image_data: Raw image data bytes
            opacity: Opacity level (0.0-1.0)

        Returns:
            Processed image data as bytes
        """
        try:
            from PIL import Image
            import io

            # Load image
            img = Image.open(io.BytesIO(image_data))

            # Convert to RGBA if not already
            if img.mode != 'RGBA':
                img = img.convert('RGBA')

            # Apply transparency to all pixels
            # Multiply existing alpha by desired opacity
            pixels = img.load()
            for y in range(img.height):
                for x in range(img.width):
                    r, g, b, a = pixels[x, y]
                    # Apply opacity to alpha channel
                    new_alpha = int(a * opacity)
                    pixels[x, y] = (r, g, b, new_alpha)

            # Save to bytes
            output = io.BytesIO()
            img.save(output, format='PNG')
            return output.getvalue()

        except Exception as e:
            logger.error(f"Failed to process image transparency: {e}")
            return image_data  # Return original if processing fails

    def apply_to_word_document(self, doc: Document, opacity: float = None) -> bool:
        """
        Apply cat background to Word document header (behind text, no layout impact)

        This method places the background image in the document header as a floating
        image using the CT_Anchor class method, ensuring it doesn't affect body content layout.

        Args:
            doc: Word document to apply background to
            opacity: Image opacity (0.0-1.0), defaults to config.BACKGROUND_TRANSPARENCY_WORD

        Returns:
            True if successful, False otherwise
        """
        try:
            # Use default opacity if not specified
            if opacity is None:
                opacity = config.BACKGROUND_TRANSPARENCY_WORD

            # Load background image if not already loaded
            if self._background_image is None:
                self._background_image = self._load_template()

            if self._background_image is None:
                logger.warning("No background image available")
                return False

            # Apply transparency using cache
            cache_key = f"{opacity:.2f}"
            if cache_key not in self._transparent_image_cache:
                self._transparent_image_cache[cache_key] = self._process_image_transparency(
                    self._background_image, opacity
                )

            background_image = self._transparent_image_cache[cache_key]

            # Apply to all sections (for multi-page documents)
            for section in doc.sections:
                # Clear any existing header content
                header = section.header
                for para in header.paragraphs:
                    para._element.getparent().remove(para._element)

                # Create new paragraph for background
                p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()

                # Add a run to hold the drawing
                run = p.add_run()
                from io import BytesIO
                img_stream = BytesIO(background_image)

                # Get page dimensions
                page_width_twips = section.page_width
                page_height_twips = section.page_height
                page_width_inches = page_width_twips / 914400
                page_height_inches = page_height_twips / 914400

                # Calculate aspect ratio
                img_aspect = self._background_width / self._background_height

                # Fit by height to preserve aspect ratio
                img_height_inches = page_height_inches
                img_width_inches = page_height_inches * img_aspect

                # Calculate horizontal offset to center (in EMUs)
                x_offset_emus = int((page_width_inches - img_width_inches) / 2 * 914400)
                y_offset_emus = 0  # Top of page

                # Create floating anchor using new_pic_anchor function
                anchor = new_pic_anchor(
                    run.part,
                    img_stream,
                    width=Inches(img_width_inches),
                    height=Inches(img_height_inches),
                    pos_x=x_offset_emus,
                    pos_y=y_offset_emus
                )

                # Add anchor to the run's drawing element
                run._r.add_drawing(anchor)

                # Set header paragraph to not take space
                pPr = OxmlElement('w:pPr')
                spacing = OxmlElement('w:spacing')
                spacing.set(qn('w:before'), '0')
                spacing.set(qn('w:after'), '0')
                spacing.set(qn('w:line'), '1')
                pPr.append(spacing)
                p._element.insert(0, pPr)

            logger.info(f"Background image applied to header (opacity: {opacity:.0%})")
            return True

        except Exception as e:
            logger.error(f"Failed to apply background to Word document: {e}")
            import traceback
            traceback.print_exc()
            return False


# CT_Anchor class for floating picture support
# Based on: https://developer.volcengine.com/articles/7383034922560618506
class CT_Anchor(BaseOxmlElement):
    """
    ``<w:anchor>`` element, container for a floating image.
    """
    extent = OneAndOnlyOne('wp:extent')
    docPr = OneAndOnlyOne('wp:docPr')
    graphic = OneAndOnlyOne('a:graphic')

    @classmethod
    def new(cls, cx, cy, shape_id, pic, pos_x, pos_y):
        """
        Return a new ``<wp:anchor>`` element populated with the values passed
        as parameters.
        """
        anchor = parse_xml(cls._anchor_xml(pos_x, pos_y))
        anchor.extent.cx = cx
        anchor.extent.cy = cy
        anchor.docPr.id = shape_id
        anchor.docPr.name = 'Picture %d' % shape_id
        anchor.graphic.graphicData.uri = (
            'http://schemas.openxmlformats.org/drawingml/2006/picture'
        )
        anchor.graphic.graphicData._insert_pic(pic)
        return anchor

    @classmethod
    def new_pic_anchor(cls, shape_id, rId, filename, cx, cy, pos_x, pos_y):
        """
        Return a new `wp:anchor` element containing the `pic:pic` element
        specified by the argument values.
        """
        pic_id = 0  # Word doesn't seem to use this, but does not omit it
        pic = CT_Picture.new(pic_id, filename, rId, cx, cy)
        anchor = cls.new(cx, cy, shape_id, pic, pos_x, pos_y)
        anchor.graphic.graphicData._insert_pic(pic)
        return anchor

    @classmethod
    def _anchor_xml(cls, pos_x, pos_y):
        return (
            '<wp:anchor distT="0" distB="0" distL="0" distR="0" simplePos="0" relativeHeight="0" \n'
            '           behindDoc="1" locked="0" layoutInCell="1" allowOverlap="1" \n'
            '           %s>\n'
            '  <wp:simplePos x="0" y="0"/>\n'
            '  <wp:positionH relativeFrom="page">\n'
            '    <wp:posOffset>%d</wp:posOffset>\n'
            '  </wp:positionH>\n'
            '  <wp:positionV relativeFrom="page">\n'
            '    <wp:posOffset>%d</wp:posOffset>\n'
            '  </wp:positionV>\n'
            '  <wp:extent cx="914400" cy="914400"/>\n'
            '  <wp:wrapNone/>\n'
            '  <wp:docPr id="666" name="unnamed"/>\n'
            '  <wp:cNvGraphicFramePr>\n'
            '    <a:graphicFrameLocks noChangeAspect="1"/>\n'
            '  </wp:cNvGraphicFramePr>\n'
            '  <a:graphic>\n'
            '    <a:graphicData uri="URI not set"/>\n'
            '  </a:graphic>\n'
            '</wp:anchor>' % (nsdecls('wp', 'a', 'pic', 'r'), int(pos_x), int(pos_y))
        )


# Register CT_Anchor for wp:anchor elements
try:
    from docx.oxml import register_element_cls
    register_element_cls('wp:anchor', CT_Anchor)
except Exception:
    # Already registered or registration not available
    pass


def new_pic_anchor(part, image_descriptor, width, height, pos_x, pos_y):
    """Return a newly-created `wp:anchor` element.

    The element contains the image specified by *image_descriptor* and is scaled
    based on the values of *width* and *height*.
    """
    rId, image = part.get_or_add_image(image_descriptor)
    cx, cy = image.scaled_dimensions(width, height)
    shape_id, filename = part.next_id, image.filename
    return CT_Anchor.new_pic_anchor(shape_id, rId, filename, cx, cy, pos_x, pos_y)


class PDFStyler:
    """
    PDF document styling utilities
    Applies professional formatting to PDF reports
    """

    @staticmethod
    def get_color(hex_color: str):
        """Convert hex color to ReportLab Color object"""
        from reportlab.lib import colors
        return colors.HexColor(hex_color)

    @staticmethod
    def setup_page_settings(doc, top_cm: float = 2.5, bottom_cm: float = 2.5,
                            left_cm: float = 2.0, right_cm: float = 2.0):
        """
        Set up PDF page margins

        Args:
            doc: SimpleDocTemplate object
            top_cm: Top margin in cm
            bottom_cm: Bottom margin in cm
            left_cm: Left margin in cm
            right_cm: Right margin in cm
        """
        from reportlab.lib.units import cm

        doc.topMargin = top_cm * cm
        doc.bottomMargin = bottom_cm * cm
        doc.leftMargin = left_cm * cm
        doc.rightMargin = right_cm * cm

    @staticmethod
    def create_style_dict() -> dict:
        """
        Create a dictionary of style configurations

        Returns:
            Dictionary with style settings
        """
        return {
            'title': {
                'font': config.FONT_MAIN_TITLE,
                'size': config.FONT_SIZE_MAIN_TITLE,
                'color': config.COLOR_ACCENT_BLUE,
                'leading': config.FONT_SIZE_MAIN_TITLE * 1.5
            },
            'heading1': {
                'font': config.FONT_HEADING1,
                'size': config.FONT_SIZE_HEADING1,
                'color': config.COLOR_ACCENT_BLUE,
                'leading': config.FONT_SIZE_HEADING1 * 1.4
            },
            'heading2': {
                'font': config.FONT_HEADING2,
                'size': config.FONT_SIZE_HEADING2,
                'color': config.COLOR_TEXT_DARK,
                'leading': config.FONT_SIZE_HEADING2 * 1.3
            },
            'body': {
                'font': config.FONT_BODY,
                'size': config.FONT_SIZE_BODY,
                'color': config.COLOR_TEXT_DARK,
                'leading': config.FONT_SIZE_BODY * config.LINE_SPACING
            },
            'data': {
                'font': config.FONT_DATA,
                'size': config.FONT_SIZE_DATA,
                'color': config.COLOR_TEXT_DARK,
                'leading': config.FONT_SIZE_DATA * config.LINE_SPACING
            }
        }

    @staticmethod
    def apply_background_to_pdf(canvas, image_path: str = None,
                                opacity: float = 0.6):
        """
        Apply background image to PDF page

        Args:
            canvas: ReportLab canvas object
            image_path: Path to background image (if None, uses default)
            opacity: Image opacity (0.0-1.0)
        """
        try:
            # If no image path provided, try to load from template
            if image_path is None:
                manager = BackgroundTemplateManager()
                image_data = manager._load_template()
                if image_data:
                    # Save to temp file
                    import tempfile
                    with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as f:
                        f.write(image_data)
                        image_path = f.name

            if image_path and os.path.exists(image_path):
                from reportlab.lib.units import inch

                # Get page dimensions
                page_width = canvas._pagesize[0]
                page_height = canvas._pagesize[1]

                # Set transparency
                canvas.saveState()
                canvas.setFillAlpha(opacity)

                # Draw image centered on page
                canvas.drawImage(
                    image_path,
                    0, 0,
                    width=page_width,
                    height=page_height,
                    preserveAspectRatio=True
                )

                canvas.restoreState()

        except Exception as e:
            logger.error(f"Failed to apply background to PDF: {e}")


# Convenience functions
def apply_professional_styling_to_word(doc: Document, use_template: bool = True,
                                        add_content_overlay: bool = True) -> Document:
    """
    Apply all professional styling to a Word document

    Args:
        doc: Word document
        use_template: Whether to apply cat background template
        add_content_overlay: Whether to add white semi-transparent overlay to content

    Returns:
        Styled document
    """
    # Set page margins
    WordStyler.set_page_margins(
        doc,
        top_cm=config.MARGIN_TOP_CM,
        bottom_cm=config.MARGIN_BOTTOM_CM,
        left_cm=config.MARGIN_LEFT_CM,
        right_cm=config.MARGIN_RIGHT_CM
    )

    # Apply background template if requested
    if use_template:
        manager = BackgroundTemplateManager()
        # Apply with configured transparency
        manager.apply_to_word_document(doc, opacity=config.BACKGROUND_TRANSPARENCY_WORD)

    # Note: Content overlay is applied during paragraph creation in report_generator.py
    # This is because Word paragraph shading needs to be applied per-paragraph

    return doc


def get_pdf_style_config() -> dict:
    """Get PDF style configuration dictionary"""
    return PDFStyler.create_style_dict()


if __name__ == "__main__":
    # Test the styling utilities
    import tempfile
    from docx.shared import Inches

    # Test Word styling
    doc = Document()

    # Apply professional styling
    apply_professional_styling_to_word(doc, use_template=False)  # Skip template for test

    # Add styled content
    WordStyler.apply_heading_style(doc, "测试标题", level=1, emoji="📊")
    WordStyler.apply_body_style(doc, "这是一个测试段落，用于验证样式是否正确应用。", indent=True)
    WordStyler.add_section_divider(doc)

    # Save test document
    test_path = os.path.join(tempfile.gettempdir(), "styling_test.docx")
    doc.save(test_path)
    print(f"✓ Styling test document saved to: {test_path}")

    # Test background template loading
    manager = BackgroundTemplateManager()
    if manager._load_template():
        print(f"✓ Background template loaded successfully: {manager._background_width}x{manager._background_height}px")
    else:
        print("✗ Background template not found")
