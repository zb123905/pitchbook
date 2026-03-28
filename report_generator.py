"""
Report Generation Module
Generate VC/PE industry weekly market observation reports based on analysis results

Enhanced with visualization charts (Phase 3)
"""
import os
import logging
from datetime import datetime
from collections import Counter
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Mm, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement
import config
from styling_utils import (
    WordStyler,
    BackgroundTemplateManager,
    apply_professional_styling_to_word
)

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)

logger = logging.getLogger(__name__)

# Phase 3 and Phase 1 availability flags (modules will be imported lazily)
VISUALIZATION_AVAILABLE = True
PDF_AVAILABLE = True

# LLM client (initialized lazily)
_llm_client = None


def get_llm_client():
    """Get or initialize LLM client for report generation"""
    global _llm_client
    if _llm_client is None:
        try:
            from llm.deepseek_client import get_default_client
            _llm_client = get_default_client()
            if _llm_client and _llm_client.is_available():
                logger.info("✓ LLM client available for report generation")
            else:
                logger.warning("LLM client not available, will use fallback mode")
                _llm_client = False  # Mark as unavailable
        except Exception as e:
            logger.warning(f"Failed to initialize LLM client: {e}")
            _llm_client = False
    return _llm_client if _llm_client is not False else None


def clean_llm_content(content: str) -> str:
    """
    Clean LLM-generated content by removing markdown code block markers

    This function removes ``` markers that sometimes appear in LLM output,
    while preserving the actual content. It handles unclosed code blocks
    by ensuring content is never lost.

    Args:
        content: Raw LLM response

    Returns:
        Cleaned content without ``` markers
    """
    import re

    # Method 1: Remove code block markers while preserving content
    # This handles both properly closed and unclosed blocks
    result = re.sub(r'```[a-zA-Z]*\n?', '', content)  # Remove opening markers
    result = re.sub(r'```\n?', '', result)  # Remove closing markers

    # Clean up any resulting empty lines
    lines = result.split('\n')
    cleaned_lines = []
    consecutive_empty = 0

    for line in lines:
        if line.strip() == '':
            consecutive_empty += 1
            if consecutive_empty <= 2:  # Keep up to 2 consecutive empty lines
                cleaned_lines.append(line)
        else:
            consecutive_empty = 0
            cleaned_lines.append(line)

    return '\n'.join(cleaned_lines).strip()


class WeeklyReportGenerator:
    """Weekly Market Observation Report Generator (Enhanced with Charts)"""

    def __init__(self, enable_charts: bool = True, use_llm: bool = True, use_template: bool = False, enable_full_article: bool = True):
        """
        Initialize report generator

        Args:
            enable_charts: Whether to include visualization charts (Phase 3 feature)
            use_llm: Whether to use LLM for content generation (default: True)
            use_template: Whether to use cat background template (default: True)
            enable_full_article: Whether to generate 3000+ word AI article (default: True)
        """
        self.enable_charts = enable_charts and VISUALIZATION_AVAILABLE
        self.use_llm = use_llm and config.ENABLE_LLM_ANALYSIS
        self.use_template = use_template
        self.enable_full_article = enable_full_article and use_llm  # Full article requires LLM
        self.visualizer = None  # Will be loaded lazily
        self.trend_analyzer = None  # Will be loaded lazily

        # Check LLM availability
        if self.use_llm:
            self.llm_client = get_llm_client()
            if self.llm_client is None:
                logger.warning("LLM requested but not available, will use fallback mode")
                self.use_llm = False
                self.enable_full_article = False  # Cannot generate article without LLM
        else:
            self.llm_client = None

        # Initialize background template manager
        # Background functionality disabled - use_template defaults to False
        # if self.use_template:
        #     self.bg_manager = BackgroundTemplateManager()

        logger.info(f"WeeklyReportGenerator initialized (charts: {self.enable_charts}, llm: {self.use_llm}, template: {self.use_template})")
        logger.info(f"Professional formatting enabled: Microsoft YaHei fonts, professional colors, {config.LINE_SPACING}x line spacing")

    def generate_weekly_report(self, analyses, output_path=None, market_overview=None, use_llm=True, use_template=None):
        """
        Generate weekly report (enhanced version)

        Args:
            analyses: List of analysis results
            output_path: Optional output file path
            market_overview: Optional market overview data
            use_llm: Whether to use LLM for content generation (default: True)
            use_template: Whether to use cat background template (default: from __init__)
        """
        # Update LLM mode
        self.use_llm = use_llm and config.ENABLE_LLM_ANALYSIS

        # Update template mode if explicitly provided
        # Background functionality disabled
        # if use_template is not None:
        #     self.use_template = use_template
        #     if self.use_template:
        #         self.bg_manager = BackgroundTemplateManager()
        if self.use_llm:
            self.llm_client = get_llm_client()
            if self.llm_client is None:
                logger.warning("LLM requested but not available, will use fallback mode")
                self.use_llm = False

        logger.info(f"Starting weekly report generation (LLM: {self.use_llm})")

        # 添加空列表检查
        if not analyses:
            logger.warning("No analyses data provided, creating minimal report")
            analyses = []

        try:
            # 创建Word文档 - 始终创建新文档（不使用模板作为基础）
            doc = Document()
            logger.info("Created new document (background will be applied via styling if enabled)")

            # Apply professional styling (includes background if use_template=True)
            apply_professional_styling_to_word(doc, use_template=self.use_template)

            # Set default font to Microsoft YaHei
            doc.styles['Normal'].font.name = config.FONT_BODY
            doc.styles['Normal'].font.size = Pt(config.FONT_SIZE_BODY)

            # Add title with professional styling
            title = doc.add_heading('VC/PE 行业每周市场观察报告', level=0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            WordStyler.set_font(title.runs[0], config.FONT_MAIN_TITLE, config.FONT_SIZE_MAIN_TITLE, bold=True, color=config.COLOR_ACCENT_BLUE)

            # Add report overview
            WordStyler.apply_heading_style(doc, '报告概览', level=1, emoji='📊')
            # WordStyler.add_section_divider(doc)  # Removed to reduce whitespace

            report_date = datetime.now().strftime('%Y年%m月%d日')
            total_items = len(analyses)
            email_analyses = [a for a in analyses if a.get('email_index') or a.get('subject')]
            report_analyses = [a for a in analyses if a.get('file_type') == 'Downloaded Report']

            # Add overview information with styling
            overview = doc.add_paragraph()
            WordStyler.set_spacing(overview, line_spacing=config.LINE_SPACING)
            overview.add_run('报告日期: ').bold = True
            overview.add_run(f'{report_date}\n')
            overview.add_run('分析项目总数: ').bold = True
            overview.add_run(f'{total_items}\n')
            overview.add_run('邮件分析: ').bold = True
            overview.add_run(f'{len(email_analyses)} 封\n')
            overview.add_run('报告分析: ').bold = True
            overview.add_run(f'{len(report_analyses)} 份\n')

            if market_overview:
                market_sentiment = market_overview.get('market_sentiment', 'neutral')
                overview.add_run('市场情绪: ').bold = True
                overview.add_run(f'{market_sentiment}\n')

            # Add executive summary with compact formatting
            WordStyler.apply_heading_style(doc, '执行摘要', level=1, emoji='📝')

            executive_summary = self._generate_executive_summary(analyses, market_overview or {})

            # Debug logging
            logger.info(f"[DEBUG] Executive summary length: {len(executive_summary)} chars, {len(executive_summary.strip())} chars (stripped)")

            # Direct approach: add content as one block with proper formatting
            if executive_summary and executive_summary.strip():
                import re
                # Clean markdown code blocks if present
                cleaned = clean_llm_content(executive_summary)

                logger.info(f"[DEBUG] Cleaned content length: {len(cleaned)} chars")

                # Add as a single paragraph with proper spacing
                p = doc.add_paragraph()
                run = p.add_run(cleaned)
                # Apply font settings
                run.font.name = config.FONT_BODY
                run.font.size = Pt(config.FONT_SIZE_BODY)
                # Compact spacing
                WordStyler.set_spacing(p, line_spacing=1.0, before=Pt(3), after=Pt(6))

                logger.info(f"[DEBUG] Executive summary paragraph added to document")
            else:
                logger.warning("[DEBUG] Executive summary is empty or None, using fallback")

            # Add full AI article (NEW - 3000+ words)
            if self.enable_full_article:
                full_article = self._generate_full_article(analyses, market_overview or {})
                if full_article:
                    # Use minimal spacing instead of page break for more compact layout
                    # doc.add_page_break()  # Removed to save space
                    # Create main heading with compact spacing
                    h = doc.add_paragraph()
                    run = h.add_run('🤖 AI深度分析报告')
                    run.bold = True
                    run.font.size = Pt(15)
                    run.font.color.rgb = RGBColor(0, 102, 204)
                    WordStyler.set_spacing(h, line_spacing=1.0, before=Pt(6), after=Pt(3))
                    # Skip section divider for more compact layout

                    # Add article intro
                    intro = doc.add_paragraph()
                    intro.add_run('以下是由AI生成的3000字深度分析报告，涵盖市场全景、投资趋势、行业洞察与投资建议。').italic = True
                    WordStyler.set_spacing(intro, line_spacing=1.0, before=Pt(2))
                    WordStyler.set_color(intro.runs[0], config.COLOR_TEXT_LIGHT)

                    # Add article content paragraph by paragraph with maximum compactness
                    for paragraph in full_article.split('\n\n'):
                        if paragraph.strip():
                            # Check if this is a heading
                            if paragraph.strip().startswith('###') or paragraph.strip().startswith('##'):
                                heading_text = paragraph.strip().lstrip('#').strip()
                                # Determine heading level - use compact style
                                if paragraph.startswith('###'):
                                    # Create compact heading manually
                                    h = doc.add_paragraph()
                                    run = h.add_run(heading_text)
                                    run.bold = True
                                    WordStyler.set_spacing(h, line_spacing=1.0, before=Pt(3), after=Pt(1))
                                else:
                                    # Level 2 heading
                                    h = doc.add_paragraph()
                                    run = h.add_run(heading_text)
                                    run.bold = True
                                    run.font.size = Pt(13)
                                    WordStyler.set_spacing(h, line_spacing=1.0, before=Pt(5), after=Pt(1))
                            else:
                                # Regular paragraph - use maximum compact formatting
                                p = doc.add_paragraph()
                                for line in paragraph.split('\n'):
                                    if line.strip():
                                        # Check for bold text markers
                                        if line.strip().startswith('**') and line.strip().endswith('**'):
                                            bold_text = line.strip().replace('**', '')
                                            run = p.add_run(bold_text + ' ')
                                            run.bold = True
                                        elif line.strip().startswith('- '):
                                            # Bullet point
                                            run = p.add_run(line.strip() + ' ')
                                        elif line.strip().startswith(('1. ', '2. ', '3. ', '4. ', '5. ', '6. ')):
                                            # Numbered list item
                                            run = p.add_run(line.strip() + ' ')
                                            run.bold = True
                                        else:
                                            run = p.add_run(line.strip() + ' ')
                                # Maximum compact spacing
                                WordStyler.set_spacing(p, line_spacing=1.0, after=Pt(0))

            # Add market overview - DISABLED to save space for page count control
            # WordStyler.apply_heading_style(doc, '市场概览', level=1, emoji='🌐')
            # WordStyler.add_section_divider(doc)  # Removed to reduce whitespace
            #
            # content_types = [a['content_type'] for a in analyses]
            # type_counts = Counter(content_types)
            #
            # WordStyler.apply_body_style(doc, '内容类型分布:', bold_prefix='')
            #
            # for content_type, count in type_counts.most_common():
            #     percentage = (count / len(analyses)) * 100
            #     p = doc.add_paragraph()
            #     p.add_run(f'{content_type}: ').bold = True
            #     p.add_run(f'{count} ({percentage:.1f}%)')
            #     WordStyler.set_spacing(p, line_spacing=config.LINE_SPACING)
            #
            # # Add industry sector analysis
            # WordStyler.apply_heading_style(doc, '行业板块分析', level=1, emoji='🏭')
            # # WordStyler.add_section_divider(doc)  # Removed to reduce whitespace
            #
            # all_topics = []
            # for analysis in analyses:
            #     all_topics.extend(analysis['key_topics'])
            #
            # if all_topics:
            #     topic_counts = Counter(all_topics)
            #
            #     WordStyler.apply_body_style(doc, '热门行业板块:', bold_prefix='')
            #
            #     for topic, count in topic_counts.most_common():
            #         percentage = (count / len(analyses)) * 100
            #         p = doc.add_paragraph()
            #         p.add_run(f'{topic}: ').bold = True
            #         p.add_run(f'{count} 次提及 ({percentage:.1f}%)')
            #         WordStyler.set_spacing(p, line_spacing=config.LINE_SPACING)

            # Add email analysis section - DISABLED to save space
            # if email_analyses:
            #     self._add_email_analysis_section(doc, email_analyses)
            #
            # # Add report analysis section (NEW!)
            # if report_analyses:
            #     self._add_report_analysis_section(doc, report_analyses)

            # Add key trends
            WordStyler.apply_heading_style(doc, '关键趋势和观察', level=1, emoji='📈')

            key_trends = self._generate_key_trends(analyses, market_overview or {})
            if key_trends.strip():
                import re
                cleaned = clean_llm_content(key_trends)
                # Add as paragraph block
                p = doc.add_paragraph()
                run = p.add_run(cleaned)
                run.font.name = config.FONT_BODY
                run.font.size = Pt(config.FONT_SIZE_BODY)
                WordStyler.set_spacing(p, line_spacing=1.0, before=Pt(3), after=Pt(6))

            # Add market recommendations
            WordStyler.apply_heading_style(doc, '市场观察和建议', level=1, emoji='💡')

            recommendations = self._generate_recommendations(analyses)
            if recommendations.strip():
                cleaned = clean_llm_content(recommendations)
                # Add as paragraph block
                p = doc.add_paragraph()
                run = p.add_run(cleaned)
                run.font.name = config.FONT_BODY
                run.font.size = Pt(config.FONT_SIZE_BODY)
                WordStyler.set_spacing(p, line_spacing=1.0, before=Pt(3), after=Pt(6))

            # Add appendix
            WordStyler.apply_heading_style(doc, '附录', level=1, emoji='📎')
            # WordStyler.add_section_divider(doc)  # Removed to reduce whitespace

            # Add charts section (Phase 3: Visualization) - DISABLED for page count control
            # if self.enable_charts:
            #     charts_section = self._add_charts_section(doc, analyses, market_overview, max_charts=2)

            WordStyler.apply_body_style(doc, '数据来源:', bold_prefix='')
            WordStyler.apply_body_style(doc, '主要数据来源: PitchBook 订阅邮件', bold_prefix='')
            WordStyler.apply_body_style(doc, f'处理周期: {datetime.now().strftime("%Y-%m-%d")}', bold_prefix='')

            # WordStyler.add_section_divider(doc)  # Removed to reduce whitespace

            WordStyler.apply_body_style(doc, '报告生成信息:', bold_prefix='')
            WordStyler.apply_body_style(doc, f'生成时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}', bold_prefix='')
            WordStyler.apply_body_style(doc, '报告版本: v4.1 (专业格式 + 猫背景)', bold_prefix='')

            # Add DeepSeek attribution if LLM was used
            if self.use_llm and self.llm_client:
                WordStyler.apply_body_style(doc, 'AI 分析: 本报告内容部分由 DeepSeek AI 生成', bold_prefix='')

            # 保存文档
            if output_path is None:
                # Add suffix for full article reports
                suffix = "_完整报告" if self.enable_full_article and full_article else ""
                output_path = os.path.join(
                    config.SUMMARY_REPORT_DIR,
                    f'VC_PE_Weekly_AI分析_{datetime.now().strftime("%Y%m%d_%H%M%S")}{suffix}.docx'
                )

            doc.save(output_path)
            logger.info(f"Weekly report saved to: {output_path}")
            return output_path

        except Exception as e:
            logger.error(f"Failed to generate weekly report: {str(e)}")
            return None

    def generate_word_report(self, analyses, market_overview):
        """Generate Word format report (legacy method for compatibility)"""
        return self.generate_weekly_report(analyses, None, market_overview)

    def _generate_executive_summary(self, analyses, market_overview):
        """Generate executive summary (with LLM support)"""
        # Try to use LLM if enabled
        if self.use_llm and self.llm_client:
            try:
                logger.info("Using LLM to generate executive summary (1/3 LLM tasks)...")
                result = self.llm_client.generate_executive_summary(analyses, "本周")

                logger.info(f"[DEBUG] LLM result success: {result['success']}, has content: {'content' in result}")
                if result.get('content'):
                    logger.info(f"[DEBUG] LLM content length: {len(result['content'])} chars")

                if result['success']:
                    logger.info("✓ Executive summary generated by LLM")
                    content = clean_llm_content(result['content'])
                    logger.info(f"[DEBUG] Cleaned content length: {len(content)} chars")
                    return content
                else:
                    logger.warning(f"LLM generation failed: {result.get('error')}, falling back to template")
            except Exception as e:
                logger.warning(f"LLM error: {e}, falling back to template")
                import traceback
                traceback.print_exc()

        # Fallback to template
        logger.info("Using template for executive summary")
        summary_lines = []

        total_emails = len(analyses)
        market_sentiment = market_overview.get('market_sentiment', 'neutral')

        # Chinese translation of sentiment
        sentiment_map = {
            'positive': '积极',
            'negative': '消极',
            'neutral': '中性'
        }
        sentiment_text = sentiment_map.get(market_sentiment, '中性')

        summary_lines.append(f"本周共分析了 {total_emails} 份PitchBook行业资讯，整体市场情绪为{sentiment_text}。")

        # Count top topics
        all_topics = []
        for analysis in analyses:
            all_topics.extend(analysis.get('key_topics', []))

        if all_topics:
            topic_counts = Counter(all_topics)
            if topic_counts:
                top_topics = topic_counts.most_common(3)
                topic_text = "、".join([topic[0] for topic in top_topics])
                summary_lines.append(f"市场热门主题包括：{topic_text}。")

        summary_lines.append("主要内容反映了当前市场关注焦点和投资方向。")

        return "\n\n".join(summary_lines)

    def _generate_key_trends(self, analyses, market_overview):
        """Generate key trends content (with LLM support)"""
        # Try to use LLM if enabled
        if self.use_llm and self.llm_client:
            try:
                logger.info("Using LLM to generate key trends (2/3 LLM tasks)...")
                result = self.llm_client.generate_key_trends(analyses, "本周")

                if result['success']:
                    logger.info("✓ Key trends generated by LLM")
                    return clean_llm_content(result['content'])
                else:
                    logger.warning(f"LLM generation failed: {result.get('error')}, falling back to template")
            except Exception as e:
                logger.warning(f"LLM error: {e}, falling back to template")

        # Fallback to template
        logger.info("Using template for key trends")
        content_lines = []

        trends_count = 0

        # Trend 1: Market sentiment
        market_sentiment = market_overview.get('market_sentiment', 'neutral')

        sentiment_map = {
            'positive': '积极',
            'negative': '消极',
            'neutral': '中性'
        }

        if market_sentiment != 'neutral':
            sentiment_text = sentiment_map.get(market_sentiment, '中性')
            trends_count += 1
            content_lines.append(f"**{trends_count}. 市场情绪{sentiment_text}**")
            confidence_desc = {
                'positive': '投资者信心持续改善',
                'negative': '投资者保持谨慎态度',
                'neutral': '市场保持稳定'
            }
            content_lines.append(f"   本周整体市场情绪为{sentiment_text}，{confidence_desc.get(market_sentiment, '市场平稳')}。")

        # Trend 2: Hot topics
        all_topics = []
        for analysis in analyses:
            all_topics.extend(analysis.get('key_topics', []))

        if all_topics:
            topic_counts = Counter(all_topics)
            if topic_counts:
                top_topic = topic_counts.most_common(1)[0][0]
                trends_count += 1
                content_lines.append(f"\n**{trends_count}. {top_topic}板块保持活跃**")
                content_lines.append(f"   {top_topic}是本周最受关注的主题，相关新闻和交易活动频繁。")

        # Trend 3: Content characteristics
        content_types = [a.get('content_type', '') for a in analyses]
        if content_types:
            dominant_type = Counter(content_types).most_common(1)[0][0]
            trends_count += 1
            content_lines.append(f"\n**{trends_count}. {dominant_type}成为主要新闻类型**")
            content_lines.append(f"   {dominant_type}在本周占据主导地位，反映了市场关键关注领域。")

        # Trend 4: Deal activities
        deal_count = sum(1 for a in analyses if 'Deal' in a.get('content_type', ''))
        if deal_count > 0:
            trends_count += 1
            content_lines.append(f"\n**{trends_count}. 交易活动{'活跃' if deal_count >= 2 else '平稳'}**")
            content_lines.append(f"   本周共有{deal_count}条交易相关新闻，市场交易活动{'活跃' if deal_count >= 2 else '相对平稳'}。")

        return "\n".join(content_lines)

    def _generate_recommendations(self, analyses):
        """Generate market recommendations content (with LLM support)"""
        # Try to use LLM if enabled
        if self.use_llm and self.llm_client:
            try:
                logger.info("Using LLM to generate recommendations (3/3 LLM tasks)...")
                result = self.llm_client.generate_recommendations(analyses, "本周")

                if result['success']:
                    logger.info("✓ Recommendations generated by LLM")
                    return clean_llm_content(result['content'])
                else:
                    logger.warning(f"LLM generation failed: {result.get('error')}, falling back to template")
            except Exception as e:
                logger.warning(f"LLM error: {e}, falling back to template")

        # Fallback to template
        logger.info("Using template for recommendations")
        content_lines = []

        content_lines.append("基于本周市场分析，提供以下观察和建议：")

        # Recommendation 1: Focus on hot sectors
        all_topics = []
        for analysis in analyses:
            all_topics.extend(analysis.get('key_topics', []))

        if all_topics:
            topic_counts = Counter(all_topics)
            if topic_counts:
                top_topics = [topic[0] for topic in topic_counts.most_common(2)]
                content_lines.append(f"\n1. **重点关注领域**")
                for topic in top_topics:
                    content_lines.append(f"   - 持续关注{topic}领域的发展动态")

        # Recommendation 2: Market strategy
        deal_count = sum(1 for a in analyses if 'Deal' in a.get('content_type', ''))
        content_lines.append(f"\n2. **市场策略建议**")
        if deal_count >= 2:
            content_lines.append("   - 市场交易活动活跃，可考虑增加投资活动")
            content_lines.append("   - 重点关注优质项目，把握投资机会")
        else:
            content_lines.append("   - 市场相对平稳，建议采取稳健投资策略")
            content_lines.append("   - 重点关注基本面良好的标的")

        # Recommendation 3: Risk warning
        content_lines.append(f"\n3. **风险提示**")
        content_lines.append("   - 注意市场波动风险，做好风险控制")
        content_lines.append("   - 关注政策变化对相关行业的影响")
        content_lines.append("   - 建议通过多元化投资降低集中度风险")

        # Recommendation 4: Continuous monitoring
        content_lines.append(f"\n4. **持续监测**")
        content_lines.append("   - 建立持续的市场监测机制")
        content_lines.append("   - 定期跟踪行业龙头企业的动态")
        content_lines.append("   - 关注新兴技术和商业模式的发展")

        return "\n".join(content_lines)

    def _generate_full_article(self, analyses, market_overview):
        """
        Generate complete 3000+ word article using LLM

        Args:
            analyses: List of analysis results
            market_overview: Market overview data

        Returns:
            Full article content as string
        """
        if not self.use_llm or not self.llm_client:
            logger.warning("LLM not available, cannot generate full article")
            return None

        try:
            logger.info("Using LLM to generate full 3000+ word article...")

            # Prepare weekly data
            weekly_data = {
                'time_range': '本周',
                'overall_sentiment': market_overview.get('market_sentiment', 'neutral'),
                'total_emails': len(analyses)
            }

            # Generate full article
            result = self.llm_client.generate_full_article(weekly_data, analyses)

            if result['success']:
                word_count = result.get('word_count', 0)
                logger.info(f"✓ Full article generated ({word_count} chars)")

                # Warn if article is too short
                if word_count < 2000:
                    logger.warning(f"Article is shorter than expected: {word_count} chars (target: 3000+)")
                elif word_count >= 3000:
                    logger.info(f"✓ Article meets 3000+ word requirement: {word_count} chars")

                return clean_llm_content(result['content'])
            else:
                logger.error(f"Failed to generate full article: {result.get('error')}")
                return None

        except Exception as e:
            logger.error(f"Error generating full article: {e}")
            import traceback
            traceback.print_exc()
            return None

    def _sentiment_chinese(self, sentiment):
        """Convert market sentiment to Chinese"""
        sentiment_map = {
            'positive': '积极',
            'negative': '消极',
            'neutral': '中性'
        }
        return sentiment_map.get(sentiment, '中性')

    def _add_email_analysis_section(self, doc, email_analyses):
        """Add email analysis section (with LLM support)"""
        # Use minimal spacing instead of page break for more compact layout
        # doc.add_page_break()  # Removed to save space
        WordStyler.apply_heading_style(doc, '邮件内容深度分析', level=1, emoji='📧')
        # WordStyler.add_section_divider(doc)  # Removed to reduce whitespace

        total_emails = min(len(email_analyses), 5)
        for idx, analysis in enumerate(email_analyses[:5], 1):  # Limit to 5 emails
            logger.info(f"Processing email {idx}/{total_emails} for report...")
            subject = analysis.get('subject', 'No Subject')
            # Truncate subject if too long
            display_subject = subject[:80] + '...' if len(subject) > 80 else subject
            WordStyler.apply_heading_style(doc, f"{idx}. {display_subject}", level=2, emoji='')

            # Email metadata with styling
            p = doc.add_paragraph()
            WordStyler.set_spacing(p, line_spacing=config.LINE_SPACING)
            p.add_run('发件人: ').bold = True
            p.add_run(f"{analysis.get('from', 'N/A')}\n")
            p.add_run('日期: ').bold = True
            p.add_run(f"{analysis.get('date', 'N/A')}\n")
            p.add_run('内容类型: ').bold = True
            p.add_run(f"{analysis.get('content_type', 'N/A')}")

            # Categories with styling
            if analysis.get('categories'):
                categories_str = " | ".join(analysis['categories'])
                WordStyler.apply_body_style(doc, f"🏷️ 分类: {categories_str}", bold_prefix='')

            # Key topics with styling
            if analysis.get('key_topics'):
                topics_str = " | ".join(analysis['key_topics'])
                WordStyler.apply_body_style(doc, f"🔑 关键主题: {topics_str}", bold_prefix='')

            # Try to get LLM deep analysis
            llm_analysis_added = False
            if self.use_llm and self.llm_client:
                try:
                    logger.info(f"Generating LLM deep analysis for email {idx}/{total_emails}...")
                    result = self.llm_client.generate_email_analysis(analysis, idx)

                    if result['success']:
                        logger.info(f"✓ Email {idx} LLM analysis generated")
                        # Add LLM analysis content
                        p = doc.add_paragraph()
                        p.add_run('📖 深度分析:').bold = True
                        WordStyler.set_spacing(p, line_spacing=1.15, after=Pt(3))

                        # Add LLM analysis content as a block
                        clean_content = clean_llm_content(result['content'])
                        if clean_content.strip():
                            p = doc.add_paragraph()
                            run = p.add_run(clean_content)
                            run.font.name = config.FONT_BODY
                            run.font.size = Pt(config.FONT_SIZE_BODY)
                            WordStyler.set_spacing(p, line_spacing=1.0, before=Pt(2), after=Pt(4))
                        llm_analysis_added = True
                    else:
                        logger.warning(f"LLM analysis failed for email {idx}")
                except Exception as e:
                    logger.warning(f"LLM error for email {idx}: {e}")

            # Fallback: Show existing LLM analysis fields if available
            if not llm_analysis_added:
                # Check for existing LLM fields
                if analysis.get('llm_analysis_full') or analysis.get('analysis_full'):
                    doc.add_paragraph()
                    p = doc.add_paragraph()
                    p.add_run('📖 深度分析:').bold = True
                    WordStyler.set_spacing(p, line_spacing=config.LINE_SPACING)
                    full_analysis = analysis.get('llm_analysis_full') or analysis.get('analysis_full', '')
                    WordStyler.apply_body_style(doc, full_analysis, bold_prefix='')

                # Show key insights
                if analysis.get('llm_key_insights') or analysis.get('key_insights'):
                    doc.add_paragraph()
                    p = doc.add_paragraph()
                    p.add_run('💡 关键洞察:').bold = True
                    WordStyler.set_spacing(p, line_spacing=config.LINE_SPACING)
                    insights = analysis.get('llm_key_insights') or analysis.get('key_insights', [])
                    for insight in insights:
                        WordStyler.apply_body_style(doc, f"  • {insight}", bold_prefix='')

                # Show summary
                if analysis.get('llm_summary_chinese') or analysis.get('summary_chinese'):
                    doc.add_paragraph()
                    p = doc.add_paragraph()
                    p.add_run('📝 内容摘要:').bold = True
                    WordStyler.set_spacing(p, line_spacing=config.LINE_SPACING)
                    summary = analysis.get('llm_summary_chinese') or analysis.get('summary_chinese', '')
                    WordStyler.apply_body_style(doc, summary, bold_prefix='')

                # Show deals if available
                if analysis.get('investment_deals'):
                    doc.add_paragraph()
                    p = doc.add_paragraph()
                    p.add_run('💰 相关交易:').bold = True
                    WordStyler.set_spacing(p, line_spacing=config.LINE_SPACING)
                    for deal in analysis.get('investment_deals', [])[:3]:
                        deal_text = f"  • {deal.get('company', '未知公司')} - {deal.get('round', '')} - {deal.get('amount', '金额未知')}"
                        if deal.get('investors'):
                            deal_text += f" (投资方: {', '.join(deal['investors'])})"
                        WordStyler.apply_body_style(doc, deal_text, bold_prefix='')

    def _add_report_analysis_section(self, doc, report_analyses):
        """Add report content analysis section (NEW!)"""
        # Use minimal spacing instead of page break for more compact layout
        # doc.add_page_break()  # Removed to save space
        WordStyler.apply_heading_style(doc, '报告内容深度分析', level=1, emoji='📊')
        # WordStyler.add_section_divider(doc)  # Removed to reduce whitespace

        WordStyler.apply_body_style(doc, f"本节对下载的 {len(report_analyses)} 份报告进行深度内容分析", bold_prefix='')

        for idx, analysis in enumerate(report_analyses, 1):
            WordStyler.apply_heading_style(doc, f"{idx}. {analysis.get('filename', 'Unknown')}", level=2, emoji='')

            # Content type with styling
            p = doc.add_paragraph()
            WordStyler.set_spacing(p, line_spacing=config.LINE_SPACING)
            p.add_run('📌 内容类型: ').bold = True
            p.add_run(f"{analysis.get('content_type', 'N/A')}")

            # Categories with styling
            if analysis.get('categories'):
                categories_str = " | ".join(analysis['categories'])
                WordStyler.apply_body_style(doc, f"🏷️ 分类: {categories_str}", bold_prefix='')

            # Key topics with styling
            if analysis.get('key_topics'):
                topics_str = " | ".join(analysis['key_topics'])
                WordStyler.apply_body_style(doc, f"🔑 关键主题: {topics_str}", bold_prefix='')

            # Content summary (first 800 characters)
            if 'full_text' in analysis and analysis['full_text']:
                summary = analysis['full_text'][:800]
                if len(analysis['full_text']) > 800:
                    summary += "..."

                doc.add_paragraph()
                p = doc.add_paragraph()
                p.add_run('📝 内容摘要:').bold = True
                WordStyler.set_spacing(p, line_spacing=config.LINE_SPACING)
                WordStyler.apply_body_style(doc, summary, bold_prefix='')

            # Metrics data with styling
            if analysis.get('metrics'):
                metrics = analysis['metrics']
                doc.add_paragraph()
                p = doc.add_paragraph()
                p.add_run('📈 统计指标:').bold = True
                WordStyler.set_spacing(p, line_spacing=config.LINE_SPACING)

                WordStyler.apply_body_style(doc, f"  • 文本长度: {metrics.get('text_length', 0):,} 字符", bold_prefix='')
                WordStyler.apply_body_style(doc, f"  • 发现金额: {metrics.get('amounts_found', 0)} 处", bold_prefix='')
                WordStyler.apply_body_style(doc, f"  • 发现百分比: {metrics.get('percentages_found', 0)} 处", bold_prefix='')

                # Show sample amounts if available
                if metrics.get('sample_amounts'):
                    WordStyler.apply_body_style(doc, f"  • 金额示例: {', '.join(metrics['sample_amounts'][:5])}", bold_prefix='')

                # Show sample percentages if available
                if metrics.get('sample_percentages'):
                    WordStyler.apply_body_style(doc, f"  • 百分比示例: {', '.join(metrics['sample_percentages'][:5])}", bold_prefix='')

    def _add_charts_section(self, doc, analyses, market_overview, max_charts=2):
        """
        Add visualization charts section to report (Phase 3 feature)

        Args:
            doc: Document object
            analyses: Analysis results
            market_overview: Market overview data
            max_charts: Maximum number of charts to add (default: 2 for page control)

        Generates and embeds charts (limited selection):
        - Industry distribution pie chart
        - Investment timeline
        - Top investors ranking
        - Deal stage distribution
        - Hot sectors bar chart
        """
        # Lazy import visualization modules to avoid circular dependency
        try:
            from visualization.visualizer import ReportVisualizer
            from visualization.trend_analyzer import MarketTrendAnalyzer
        except ImportError as e:
            logger.warning(f"Cannot import visualization modules: {e}")
            doc.add_paragraph('⚠️ 可视化模块不可用 (Visualization modules not available)')
            return None

        try:
            # Use minimal spacing instead of page break for more compact layout
            WordStyler.apply_compact_heading_style(doc, '数据可视化分析 (Data Visualization)', level=1, emoji='📊')

            WordStyler.apply_body_style(
                doc,
                '本节包含基于分析数据生成的可视化图表，'
                '提供投资趋势、行业分布、投资机构活跃度等关键指标的直观展示。',
                bold_prefix=''
            )

            # Generate all charts
            output_dir = config.CHART_TEMP_DIR
            os.makedirs(output_dir, exist_ok=True)

            # Create visualizer
            visualizer = ReportVisualizer()
            charts = visualizer.create_dashboard(analyses, output_dir)

            if not charts:
                doc.add_paragraph('暂无足够数据生成图表。')
                return None

            # Define chart priority (most important first for page control)
            chart_priority = [
                ('industry_distribution', '行业分布 (Industry Distribution)', '行业板块分布饼图'),
                ('hot_sectors_bar', '热门投资领域 (Hot Investment Sectors)', '热门领域投资金额排名'),
                ('deal_stage_pie', '融资轮次分布 (Deal Stage Distribution)', '各融资轮次占比'),
                ('top_investors', '活跃投资机构 (Top Investors)', '交易数量排名'),
            ]

            # Add charts to document (limited by max_charts)
            chart_count = 0
            for chart_key, heading_text, caption in chart_priority:
                if chart_count >= max_charts:
                    break
                if charts.get(chart_key):
                    WordStyler.apply_compact_heading_style(doc, heading_text, level=2, emoji='')
                    self._add_chart_to_doc(doc, charts[chart_key], caption, compact=True)
                    chart_count += 1

            logger.info(f"Added {chart_count} charts to report (limited to {max_charts} for page control)")
            return chart_count

        except Exception as e:
            logger.error(f"Failed to add charts section: {e}")
            doc.add_paragraph(f'图表生成失败: {str(e)}')
            return None

    def _add_chart_to_doc(self, doc, image_path, caption=None, compact=False):
        """
        Add chart image to Word document

        Args:
            doc: Document object
            image_path: Path to chart image file
            caption: Optional caption text
            compact: Whether to use smaller image size for page control (default: False)
        """
        if not os.path.exists(image_path):
            logger.warning(f"Chart image not found: {image_path}")
            return

        try:
            # Add image with appropriate size (smaller if compact mode)
            width = Inches(4.5) if compact else Inches(6.0)
            doc.add_picture(image_path, width=width)

            # Add caption if provided
            if caption:
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(f'图: {caption}')
                WordStyler.set_font(run, config.FONT_DATA, config.FONT_SIZE_DATA, color=config.COLOR_TEXT_DARK)
                # Compact spacing for caption
                WordStyler.set_spacing(p, line_spacing=1.0, after=Pt(2))

        except Exception as e:
            logger.error(f"Failed to add chart to document: {e}")
