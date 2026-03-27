"""
Unit tests for styling utilities module
Tests Word and PDF styling functionality
"""
import os
import sys
import unittest
import tempfile
from unittest.mock import Mock, patch, MagicMock

# Add parent directory to path
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))

from styling_utils import (
    WordStyler,
    BackgroundTemplateManager,
    PDFStyler,
    hex_to_rgb,
    apply_professional_styling_to_word,
    get_pdf_style_config
)
from docx import Document
from docx.shared import Pt, Inches
import config


class TestHexToRgb(unittest.TestCase):
    """Test hex to RGB color conversion"""

    def test_hex_to_rgb_basic(self):
        """Test basic hex color conversion"""
        color = hex_to_rgb('#0066CC')
        self.assertEqual(list(color), [0, 102, 204])

    def test_hex_to_rgb_without_hash(self):
        """Test hex color without hash prefix"""
        color = hex_to_rgb('FF0000')
        self.assertEqual(list(color), [255, 0, 0])

    def test_hex_to_rgb_gray(self):
        """Test gray color conversion"""
        color = hex_to_rgb('#333333')
        self.assertEqual(list(color), [51, 51, 51])


class TestWordStyler(unittest.TestCase):
    """Test Word styling utilities"""

    def setUp(self):
        """Set up test document"""
        self.doc = Document()

    def test_set_font(self):
        """Test font setting"""
        p = self.doc.add_paragraph()
        run = p.add_run("Test text")

        WordStyler.set_font(run, '微软雅黑', 12, bold=True, color='#0066CC')

        self.assertEqual(run.font.name, '微软雅黑')
        self.assertEqual(run.font.size.pt, 12)
        self.assertTrue(run.font.bold)

    def test_set_color(self):
        """Test color setting"""
        p = self.doc.add_paragraph()
        run = p.add_run("Test text")

        WordStyler.set_color(run, '#FF0000')
        self.assertEqual(list(run.font.color.rgb), [255, 0, 0])

    def test_set_spacing(self):
        """Test spacing setting"""
        p = self.doc.add_paragraph("Test")

        WordStyler.set_spacing(p, line_spacing=1.5, before=6, after=6)

        self.assertEqual(p.paragraph_format.line_spacing, 1.5)

    def test_set_first_line_indent(self):
        """Test first line indent"""
        p = self.doc.add_paragraph("Test paragraph with indent")

        WordStyler.set_first_line_indent(p, chars=2)

        # Check that indent is set (value may vary by font)
        self.assertIsNotNone(p.paragraph_format.first_line_indent)

    def test_add_section_divider(self):
        """Test section divider"""
        WordStyler.add_section_divider(self.doc, color='#EEEEEE', thickness_pt=0.5)

        # Check that a paragraph was added
        self.assertGreater(len(self.doc.paragraphs), 0)

    def test_apply_heading_style(self):
        """Test heading style application"""
        WordStyler.apply_heading_style(self.doc, '测试标题', level=1, emoji='📊')

        # Check that heading was added
        self.assertGreater(len(self.doc.paragraphs), 0)
        heading = self.doc.paragraphs[0]
        self.assertIn('测试标题', heading.text)

    def test_apply_body_style(self):
        """Test body paragraph style"""
        WordStyler.apply_body_style(self.doc, '这是正文内容。', indent=True)

        # Check that paragraph was added
        self.assertGreater(len(self.doc.paragraphs), 0)

    def test_set_page_margins(self):
        """Test page margin setting"""
        WordStyler.set_page_margins(
            self.doc,
            top_cm=2.5,
            bottom_cm=2.5,
            left_cm=2.0,
            right_cm=2.0
        )

        # Check that margins are set
        section = self.doc.sections[0]
        # Convert to cm for comparison (approximate)
        self.assertAlmostEqual(section.top_margin.cm, 2.5, places=1)


class TestBackgroundTemplateManager(unittest.TestCase):
    """Test background template management"""

    def setUp(self):
        """Set up test manager"""
        self.manager = BackgroundTemplateManager()

    @patch('os.path.exists')
    def test_load_template_missing_file(self, mock_exists):
        """Test loading non-existent template"""
        mock_exists.return_value = False

        result = self.manager._load_template()

        self.assertIsNone(result)

    @patch('os.path.exists')
    def test_apply_to_word_document_no_template(self, mock_exists):
        """Test applying background when template doesn't exist"""
        mock_exists.return_value = False

        doc = Document()
        result = self.manager.apply_to_word_document(doc)

        self.assertFalse(result)


class TestPDFStyler(unittest.TestCase):
    """Test PDF styling utilities"""

    @unittest.skipIf(not os.path.exists(os.path.join(os.path.dirname(__file__), '../pdf/pdf_report_generator.py')),
                     "PDF module not available")
    def test_get_color(self):
        """Test color conversion for PDF"""
        try:
            color = PDFStyler.get_color('#0066CC')

            from reportlab.lib import colors
            self.assertIsInstance(color, colors.HexColor)
        except ImportError:
            self.skipTest("ReportLab not installed")

    def test_create_style_dict(self):
        """Test style dictionary creation"""
        style_dict = PDFStyler.create_style_dict()

        self.assertIn('title', style_dict)
        self.assertIn('heading1', style_dict)
        self.assertIn('heading2', style_dict)
        self.assertIn('body', style_dict)
        self.assertIn('data', style_dict)

        # Check that title style has correct properties
        title_style = style_dict['title']
        self.assertEqual(title_style['font'], config.FONT_MAIN_TITLE)
        self.assertEqual(title_style['size'], config.FONT_SIZE_MAIN_TITLE)
        self.assertEqual(title_style['color'], config.COLOR_ACCENT_BLUE)


class TestProfessionalStyling(unittest.TestCase):
    """Test professional styling application"""

    def test_apply_professional_styling_to_word(self):
        """Test applying all professional styling"""
        doc = Document()

        styled_doc = apply_professional_styling_to_word(doc, use_template=False)

        # Check that margins are set
        section = styled_doc.sections[0]
        self.assertAlmostEqual(section.top_margin.cm, config.MARGIN_TOP_CM, places=1)
        self.assertAlmostEqual(section.left_margin.cm, config.MARGIN_LEFT_CM, places=1)

    def test_get_pdf_style_config(self):
        """Test getting PDF style configuration"""
        style_config = get_pdf_style_config()

        self.assertIsInstance(style_config, dict)
        self.assertIn('title', style_config)
        self.assertIn('body', style_config)


class TestConfigConstants(unittest.TestCase):
    """Test that config constants are properly defined"""

    def test_margin_constants(self):
        """Test margin constants"""
        self.assertTrue(hasattr(config, 'MARGIN_TOP_CM'))
        self.assertTrue(hasattr(config, 'MARGIN_BOTTOM_CM'))
        self.assertTrue(hasattr(config, 'MARGIN_LEFT_CM'))
        self.assertTrue(hasattr(config, 'MARGIN_RIGHT_CM'))

        self.assertEqual(config.MARGIN_TOP_CM, 2.5)
        self.assertEqual(config.MARGIN_BOTTOM_CM, 2.5)
        self.assertEqual(config.MARGIN_LEFT_CM, 2.0)
        self.assertEqual(config.MARGIN_RIGHT_CM, 2.0)

    def test_font_constants(self):
        """Test font constants"""
        self.assertTrue(hasattr(config, 'FONT_MAIN_TITLE'))
        self.assertTrue(hasattr(config, 'FONT_HEADING1'))
        self.assertTrue(hasattr(config, 'FONT_HEADING2'))
        self.assertTrue(hasattr(config, 'FONT_BODY'))
        self.assertTrue(hasattr(config, 'FONT_DATA'))

        self.assertEqual(config.FONT_MAIN_TITLE, '微软雅黑')
        self.assertEqual(config.FONT_BODY, '微软雅黑')

    def test_font_size_constants(self):
        """Test font size constants"""
        self.assertTrue(hasattr(config, 'FONT_SIZE_MAIN_TITLE'))
        self.assertTrue(hasattr(config, 'FONT_SIZE_HEADING1'))
        self.assertTrue(hasattr(config, 'FONT_SIZE_HEADING2'))
        self.assertTrue(hasattr(config, 'FONT_SIZE_BODY'))
        self.assertTrue(hasattr(config, 'FONT_SIZE_DATA'))

        self.assertEqual(config.FONT_SIZE_MAIN_TITLE, 22)
        self.assertEqual(config.FONT_SIZE_HEADING1, 15)
        self.assertEqual(config.FONT_SIZE_HEADING2, 14)
        self.assertEqual(config.FONT_SIZE_BODY, 12)
        self.assertEqual(config.FONT_SIZE_DATA, 10.5)

    def test_color_constants(self):
        """Test color constants"""
        self.assertTrue(hasattr(config, 'COLOR_TEXT_DARK'))
        self.assertTrue(hasattr(config, 'COLOR_ACCENT_BLUE'))
        self.assertTrue(hasattr(config, 'COLOR_DIVIDER'))
        self.assertTrue(hasattr(config, 'COLOR_HIGHLIGHT_BG'))
        self.assertTrue(hasattr(config, 'COLOR_BORDER_LIGHT'))

        self.assertEqual(config.COLOR_TEXT_DARK, '#333333')
        self.assertEqual(config.COLOR_ACCENT_BLUE, '#0066CC')
        self.assertEqual(config.COLOR_DIVIDER, '#EEEEEE')

    def test_spacing_constants(self):
        """Test spacing constants"""
        self.assertTrue(hasattr(config, 'LINE_SPACING'))
        self.assertTrue(hasattr(config, 'SPACING_BEFORE_PARA'))
        self.assertTrue(hasattr(config, 'SPACING_AFTER_PARA'))
        self.assertTrue(hasattr(config, 'FIRST_LINE_INDENT_CHARS'))

        self.assertEqual(config.LINE_SPACING, 1.5)
        self.assertEqual(config.FIRST_LINE_INDENT_CHARS, 2)


class TestIntegration(unittest.TestCase):
    """Integration tests for styling"""

    def test_create_styled_document(self):
        """Test creating a fully styled document"""
        doc = Document()

        # Apply professional styling
        apply_professional_styling_to_word(doc, use_template=False)

        # Add content
        WordStyler.apply_heading_style(doc, '报告标题', level=1, emoji='📊')
        WordStyler.apply_body_style(doc, '这是正文内容。', indent=True)
        WordStyler.add_section_divider(doc)

        # Save to temp file
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
            temp_path = f.name

        try:
            doc.save(temp_path)

            # Check file exists and has content
            self.assertTrue(os.path.exists(temp_path))
            self.assertGreater(os.path.getsize(temp_path), 0)

        finally:
            # Clean up
            if os.path.exists(temp_path):
                os.remove(temp_path)


if __name__ == '__main__':
    # Run tests
    unittest.main(verbosity=2)
